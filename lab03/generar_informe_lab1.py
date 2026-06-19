"""
generar_informe_lab1.py
Genera el informe APA completo de la Práctica de Laboratorio 1
para el curso Diseño de Proyectos de Innovación — TECSUP 2026-I
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ─────────────────────────────────────────────────────────────
# ESTILOS GLOBALES
# ─────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21.59)   # Carta
section.page_height = Cm(27.94)
section.left_margin   = Cm(2.54)
section.right_margin  = Cm(2.54)
section.top_margin    = Cm(2.54)
section.bottom_margin = Cm(2.54)

style_normal = doc.styles['Normal']
style_normal.font.name = 'Times New Roman'
style_normal.font.size = Pt(12)
pf = style_normal.paragraph_format
pf.line_spacing = Pt(24)   # doble espacio ≈ 24 pt
pf.space_after  = Pt(0)
pf.first_line_indent = Cm(1.27)

# Heading 1
h1 = doc.styles['Heading 1']
h1.font.name = 'Times New Roman'
h1.font.size = Pt(14)
h1.font.bold = True
h1.font.color.rgb = RGBColor(0, 0, 0)
h1.paragraph_format.space_before      = Pt(24)
h1.paragraph_format.space_after       = Pt(12)
h1.paragraph_format.alignment         = WD_ALIGN_PARAGRAPH.CENTER
h1.paragraph_format.page_break_before = False   # evitar página en blanco doble

# Heading 2
h2 = doc.styles['Heading 2']
h2.font.name = 'Times New Roman'
h2.font.size = Pt(12)
h2.font.bold = True
h2.font.color.rgb = RGBColor(0, 0, 0)
h2.paragraph_format.space_before      = Pt(18)
h2.paragraph_format.space_after       = Pt(6)
h2.paragraph_format.first_line_indent = Cm(1.27)
h2.paragraph_format.alignment         = WD_ALIGN_PARAGRAPH.LEFT
h2.paragraph_format.page_break_before = False

# Heading 3
h3 = doc.styles['Heading 3']
h3.font.name = 'Times New Roman'
h3.font.size = Pt(12)
h3.font.bold = True
h3.font.italic = True
h3.font.color.rgb = RGBColor(0, 0, 0)
h3.paragraph_format.space_before      = Pt(12)
h3.paragraph_format.space_after       = Pt(0)
h3.paragraph_format.first_line_indent = Cm(1.27)
h3.paragraph_format.alignment         = WD_ALIGN_PARAGRAPH.LEFT
h3.paragraph_format.page_break_before = False


def p(text='', bold=False, italic=False, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
      space_before=0, indent=True, color=None):
    """Agrega párrafo con estilo APA."""
    par = doc.add_paragraph()
    par.alignment = align
    par.paragraph_format.line_spacing = Pt(24)
    par.paragraph_format.space_before = Pt(space_before)
    par.paragraph_format.space_after  = Pt(0)
    if indent:
        par.paragraph_format.first_line_indent = Cm(1.27)
    else:
        par.paragraph_format.first_line_indent = Pt(0)
    run = par.add_run(text)
    run.font.name  = 'Times New Roman'
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    return par


def heading(text, level=1):
    return doc.add_heading(text, level=level)


def blank():
    p('')


def page_break():
    doc.add_page_break()


def add_table(headers, rows, caption=None, note=None, caption_num=None):
    """Tabla estilo APA."""
    if caption:
        cap_par = doc.add_paragraph()
        cap_par.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cap_par.paragraph_format.space_before = Pt(18)
        cap_par.paragraph_format.space_after  = Pt(0)
        cap_par.paragraph_format.first_line_indent = Pt(0)
        r1 = cap_par.add_run(f'Tabla {caption_num}' if caption_num else 'Tabla')
        r1.font.name = 'Times New Roman'; r1.font.size = Pt(12); r1.font.bold = True; r1.font.italic = False
        cap_par2 = doc.add_paragraph()
        cap_par2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cap_par2.paragraph_format.space_before = Pt(0)
        cap_par2.paragraph_format.space_after  = Pt(2)
        cap_par2.paragraph_format.first_line_indent = Pt(0)
        r2 = cap_par2.add_run(caption)
        r2.font.name = 'Times New Roman'; r2.font.size = Pt(12); r2.font.italic = True

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cp = cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(2)
        cp.paragraph_format.space_after  = Pt(2)
        run = cp.add_run(h)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.bold = True
        # top border
        tcPr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        for side in ['top']:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '12')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            borders.append(border)
        tcPr.append(borders)

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cp = cell.paragraphs[0]
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_before = Pt(2)
            cp.paragraph_format.space_after  = Pt(2)
            run = cp.add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)

    # Bottom border on last row
    last_row = table.rows[-1]
    for cell in last_row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        for side in ['bottom']:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '12')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            borders.append(border)
        tcPr.append(borders)

    if note:
        note_par = doc.add_paragraph()
        note_par.alignment = WD_ALIGN_PARAGRAPH.LEFT
        note_par.paragraph_format.space_before = Pt(2)
        note_par.paragraph_format.space_after  = Pt(12)
        note_par.paragraph_format.first_line_indent = Pt(0)
        run_n = note_par.add_run('Nota. ')
        run_n.font.name = 'Times New Roman'; run_n.font.size = Pt(10); run_n.font.italic = True; run_n.font.bold = False
        run_v = note_par.add_run(note)
        run_v.font.name = 'Times New Roman'; run_v.font.size = Pt(10)

    return table


# ─────────────────────────────────────────────────────────────
# PÁGINA 1 — PORTADA
# ─────────────────────────────────────────────────────────────
def portada():
    inst = doc.add_paragraph()
    inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    inst.paragraph_format.space_before = Pt(72)
    inst.paragraph_format.space_after  = Pt(0)
    inst.paragraph_format.first_line_indent = Pt(0)
    r = inst.add_run('TECSUP')
    r.font.name = 'Times New Roman'; r.font.size = Pt(14); r.font.bold = True

    p('Departamento Académico de Tecnología Digital', align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    p('Carrera Profesional de Diseño y Desarrollo de Software', align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)

    blank(); blank(); blank()

    title_par = doc.add_paragraph()
    title_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_par.paragraph_format.first_line_indent = Pt(0)
    tr = title_par.add_run(
        'Laboratorio – Práctica 1\n'
        'Análisis de Mercado y Validación de Demanda:\n'
        'LegendaryClass — Sistema SaaS de Gamificación Educativa'
    )
    tr.font.name = 'Times New Roman'; tr.font.size = Pt(14); tr.font.bold = True

    blank(); blank(); blank()

    p('Curso: Diseño de Proyectos de Innovación', align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    p('Docente: Turkowsky Vizcarra, Luisa', align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    blank()
    p('Integrantes:', align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, indent=False)
    for alumno in [
        'Aguirre Saavedra, Juan Alexis',
        'Alfonso Solorzano, Samir Haziel',
        'Galvan Morales, Luis Enrique',
        'Galvan Guerrero, Matias',
    ]:
        p(alumno, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)

    blank(); blank()
    p('Lima, Perú', align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    p('2026', align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)

portada()
page_break()

# ─────────────────────────────────────────────────────────────
# ÍNDICE DE CONTENIDOS
# ─────────────────────────────────────────────────────────────
heading('Índice de Contenidos', 1)

toc_items = [
    ('Introducción', '4'),
    ('Objetivos', '6'),
    ('Objetivo general', '6'),
    ('Objetivos específicos', '6'),
    ('Capítulo I: Estudio de Mercado y Viabilidad del Proyecto', '7'),
    ('1.1  Equipo innovador y habilidades', '7'),
    ('1.2  Motivación y problemática', '8'),
    ('1.3  Ideas previas exploradas', '9'),
    ('1.4  Descripción del producto LegendaryClass', '9'),
    ('1.5  Elevator Pitch', '10'),
    ('1.6  Análisis PESTEL', '11'),
    ('1.7  Cinco fuerzas de Porter', '13'),
    ('1.8  Estrategias derivadas del análisis PESTEL y Porter', '14'),
    ('1.9  Justificación del producto', '15'),
    ('1.10 Segmentación de mercado', '15'),
    ('1.11 Análisis de la encuesta y fórmula muestral', '17'),
    ('1.12 Validación del mercado', '20'),
    ('1.13 Estimación de la demanda e ingresos', '21'),
    ('1.14 Proyección financiera a cinco años', '22'),
    ('Conclusiones', '24'),
    ('Recomendaciones', '25'),
    ('Referencias', '26'),
]
for item, pg in toc_items:
    toc_p = doc.add_paragraph()
    toc_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    toc_p.paragraph_format.first_line_indent = Pt(0)
    toc_p.paragraph_format.space_before = Pt(0)
    toc_p.paragraph_format.space_after  = Pt(0)
    toc_p.paragraph_format.line_spacing = Pt(20)
    tab_stops = toc_p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(14.5))
    r_item = toc_p.add_run(item + '\t' + pg)
    r_item.font.name = 'Times New Roman'
    r_item.font.size = Pt(12)

blank()
heading('Índice de Tablas', 1)
tabla_items = [
    ('Tabla 1   Perfil del equipo innovador', '7'),
    ('Tabla 2   Análisis PESTEL', '11'),
    ('Tabla 3   Cinco fuerzas de Porter', '13'),
    ('Tabla 4   Segmentación de mercado en cascada', '16'),
    ('Tabla 5   Distribución de respuestas — P7 (motivación)', '18'),
    ('Tabla 6   Distribución de respuestas — P13 (disposición a pagar)', '18'),
    ('Tabla 7   Cálculo del precio ponderado mensual', '19'),
    ('Tabla 8   Escenarios de demanda — Año 1', '21'),
    ('Tabla 9   Proyección financiera a cinco años (escenario moderado)', '22'),
]
for item, pg in tabla_items:
    ti_p = doc.add_paragraph()
    ti_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    ti_p.paragraph_format.first_line_indent = Pt(0)
    ti_p.paragraph_format.space_before = Pt(0)
    ti_p.paragraph_format.space_after  = Pt(0)
    ti_p.paragraph_format.line_spacing = Pt(20)
    r_ti = ti_p.add_run(item + '\t' + pg)
    r_ti.font.name = 'Times New Roman'
    r_ti.font.size = Pt(12)

blank()
heading('Índice de Figuras', 1)
fig_items = [
    ('Figura 1   Logo y pantalla principal de LegendaryClass', '10'),
    ('Figura 2   Escenarios de demanda — Año 1 (gráfico de barras)', '21'),
    ('Figura 3   Proyección financiera 2027–2031 (gráfico de líneas)', '23'),
]
for item, pg in fig_items:
    fi_p = doc.add_paragraph()
    fi_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fi_p.paragraph_format.first_line_indent = Pt(0)
    fi_p.paragraph_format.space_before = Pt(0)
    fi_p.paragraph_format.space_after  = Pt(0)
    fi_p.paragraph_format.line_spacing = Pt(20)
    r_fi = fi_p.add_run(item + '\t' + pg)
    r_fi.font.name = 'Times New Roman'
    r_fi.font.size = Pt(12)

page_break()

# ─────────────────────────────────────────────────────────────
# INTRODUCCIÓN
# ─────────────────────────────────────────────────────────────
heading('Introducción', 1)

p('La educación peruana enfrenta un desafío estructural relacionado con la motivación y el compromiso de los estudiantes en el aula. Datos del Ministerio de Educación del Perú (MINEDU, 2024) revelan que aproximadamente el 58 % de los docentes de educación básica regular reporta niveles bajos o muy bajos de participación activa de sus alumnos durante las sesiones de aprendizaje, fenómeno que se agudiza en secundaria y en los primeros grados de educación superior técnica.')

p('En este contexto, la gamificación —definida como la aplicación de mecánicas de juego en entornos no lúdicos (Deterding et al., 2011)— se consolida como una estrategia pedagógica con evidencia empírica de mejora en la motivación intrínseca, la retención de conocimiento y la asistencia (Hamari et al., 2014; Plass et al., 2015). Sin embargo, las herramientas disponibles en el mercado (Kahoot!, ClassDojo, Duolingo for Schools) están diseñadas para contextos anglosajones, no contemplan el currículo nacional peruano y tienen precios de suscripción que superan la capacidad de pago individual de la mayoría de docentes peruanos.')

p('LegendaryClass surge como respuesta a esta brecha. Es un sistema SaaS (Software as a Service) de gamificación educativa que transforma cada aula en una experiencia de rol (RPG): los estudiantes eligen un personaje, acumulan puntos de experiencia (XP) por sus logros académicos y de comportamiento, suben de nivel, completan misiones y canjean recompensas definidas por su docente. Los profesores disponen de un panel de gestión para configurar comportamientos, misiones y recompensas por aula, y los directivos pueden monitorear el desempeño institucional en tiempo real.')

p('El presente informe documenta el análisis de mercado, la validación de la demanda y la proyección financiera del proyecto, en cumplimiento de los requisitos de la Práctica de Laboratorio 1 del curso Diseño de Proyectos de Innovación (TECSUP, 2026-I). Para ello se diseñó e implementó una encuesta de preferencia de mercado aplicada a una muestra representativa de docentes peruanos, cuyos resultados permiten estimar el mercado objetivo, el precio de adopción y los ingresos proyectados a cinco años bajo escenarios conservadores, moderados y optimistas.')

page_break()

# ─────────────────────────────────────────────────────────────
# OBJETIVOS
# ─────────────────────────────────────────────────────────────
heading('Objetivos', 1)

heading('Objetivo General', 2)
p('Determinar la viabilidad comercial de LegendaryClass como solución SaaS de gamificación educativa en el mercado peruano de educación básica regular y superior técnica, mediante un análisis de mercado primario y secundario que establezca el tamaño del mercado objetivo, el precio ponderado de adopción y los ingresos esperados en un horizonte de cinco años.')

heading('Objetivos Específicos', 2)
obj_list = [
    'Identificar el universo de docentes activos en el Perú y segmentarlo en cascada (potencial, disponible, efectivo y objetivo) aplicando criterios cuantitativos obtenidos de la encuesta de preferencia.',
    'Diseñar y aplicar un instrumento de recolección de datos primarios (encuesta de 17 ítems, n = 100) que permita estimar la disposición a pagar de los docentes y la prevalencia del problema de baja motivación estudiantil.',
    'Calcular el precio mensual ponderado a partir de las respuestas de P13 y compararlo con benchmarks de mercado (Kahoot! Pro, ClassDojo Plus).',
    'Proyectar los ingresos anuales del proyecto bajo tres escenarios (optimista, moderado, pesimista) durante los ejercicios 2026–2030.',
    'Analizar el entorno macroeconómico (PESTEL) y competitivo (cinco fuerzas de Porter) para identificar factores críticos de éxito y riesgos del negocio.',
]
for i, obj in enumerate(obj_list, 1):
    obj_p = doc.add_paragraph(style='List Number')
    obj_p.paragraph_format.first_line_indent = Pt(0)
    obj_p.paragraph_format.left_indent = Cm(1.27)
    obj_p.paragraph_format.space_after = Pt(0)
    obj_p.paragraph_format.line_spacing = Pt(24)
    run = obj_p.runs[0] if obj_p.runs else obj_p.add_run('')
    # Limpiar y añadir texto
    for run in obj_p.runs:
        run.text = ''
    r = obj_p.add_run(obj)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

page_break()

# ─────────────────────────────────────────────────────────────
# CAPÍTULO I
# ─────────────────────────────────────────────────────────────
heading('Capítulo I: Estudio de Mercado y Viabilidad del Proyecto', 1)

# 1.1 Equipo
heading('1.1 Equipo Innovador y Habilidades', 2)
p('El equipo LegendaryClass está conformado por cuatro estudiantes de la carrera Diseño y Desarrollo de Software de TECSUP, con perfiles complementarios que abarcan las dimensiones técnicas, de diseño y de gestión de proyectos requeridas para el desarrollo de un producto SaaS educativo.')

add_table(
    headers=['Integrante', 'Rol principal', 'Habilidades clave'],
    rows=[
        ['Aguirre Saavedra, Juan Alexis', 'Líder técnico / Full-stack', 'Angular 18, NestJS, Prisma ORM, PostgreSQL, diseño de arquitecturas SaaS'],
        ['Alfonso Solorzano, Samir Haziel', 'Backend / Base de datos', 'Node.js, REST API, modelado de datos, integración de pagos'],
        ['Galvan Morales, Luis Enrique', 'Diseño UX/UI', 'Figma, Tailwind CSS, accesibilidad web, prototipado de alta fidelidad'],
        ['Galvan Guerrero, Matias', 'QA / Analista de negocio', 'Testing funcional, análisis de requisitos, documentación técnica'],
    ],
    caption='Perfil del equipo innovador LegendaryClass',
    note='Elaboración propia (2026).',
    caption_num=1,
)

# 1.2 Motivación
heading('1.2 Motivación y Problemática', 2)
p('La motivación para desarrollar LegendaryClass surgió de la experiencia directa de los integrantes del equipo como estudiantes y, en algunos casos, como ayudantes de cátedra en talleres de programación. Se observó de manera recurrente que los estudiantes muestran mayor compromiso cuando las actividades incorporan elementos lúdicos: tablas de clasificación, insignias de logro, retos cronometrados y sistemas de recompensa.')
p('Esta observación coincide con la literatura especializada. Hamari et al. (2014), en su metaanálisis de 24 estudios sobre gamificación en educación, concluyen que el uso de mecánicas de juego produce efectos positivos estadísticamente significativos en la motivación intrínseca y la participación activa, especialmente cuando el diseño es coherente con los objetivos de aprendizaje.')
p('El problema central es la ausencia en el mercado latinoamericano de una plataforma de gamificación educativa asequible, adaptada al contexto cultural peruano, integrada con el currículo nacional y diseñada específicamente para docentes de educación básica y superior técnica. Las soluciones existentes presentan tres barreras principales: (a) precios en dólares que superan la capacidad de pago docente, (b) interfaz en inglés sin soporte local y (c) mecánicas de juego genéricas no alineadas con las dinámicas del aula peruana.')

# 1.3 Ideas previas
heading('1.3 Ideas Previas Exploradas', 2)
p('Antes de converger en LegendaryClass, el equipo evaluó tres conceptos alternativos:')
ideas = [
    ('App de seguimiento de tareas gamificado', 'Se descartó porque ya existe saturación en el mercado (Habitica, Todoist) y no resuelve la dinámica grupal del aula.'),
    ('Sistema de asistencia con reconocimiento facial', 'Descartado por el alto costo de infraestructura, barreras regulatorias de protección de datos y resistencia docente al monitoreo biométrico.'),
    ('Marketplace de recursos didácticos', 'El modelo de negocio requería masa crítica de contenidos antes de generar ingresos; el tiempo al mercado era incompatible con el horizonte del proyecto.'),
]
for idea, razon in ideas:
    ip = doc.add_paragraph()
    ip.paragraph_format.first_line_indent = Cm(1.27)
    ip.paragraph_format.line_spacing = Pt(24)
    ip.paragraph_format.space_after = Pt(0)
    r1 = ip.add_run(idea + '. ')
    r1.font.name = 'Times New Roman'; r1.font.size = Pt(12); r1.font.bold = True; r1.font.italic = True
    r2 = ip.add_run(razon)
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(12)

p('LegendaryClass fue seleccionado porque combina un problema validado (baja motivación), una solución diferenciada (RPG educativo contextualizado), un modelo de negocio probado (SaaS por suscripción) y un mercado objetivo grande y accesible (docentes peruanos con acceso a internet).')

# 1.4 Descripción del producto
heading('1.4 Descripción del Producto: LegendaryClass', 2)
p('LegendaryClass es una plataforma web SaaS de gamificación educativa diseñada para docentes y estudiantes de educación básica regular (niveles primaria y secundaria) y educación superior técnica en el Perú. La propuesta de valor central es transformar el aula en una experiencia de rol (RPG) que incrementa la motivación intrínseca y la participación activa de los estudiantes.')
p('Los módulos principales son: (a) Gestión de aulas, donde el docente crea sus grupos, genera un código de acceso y configura comportamientos, misiones y recompensas; (b) Perfil del aventurero, donde cada estudiante elige un personaje (Mago, Guerrero, Ninja, Arquero o Lanzador), acumula XP y sube de nivel; (c) Sistema de puntos y comportamientos, que permite al docente otorgar o descontar puntos en tiempo real durante la clase; (d) Misiones y retos, tareas académicas configurables con fecha límite y recompensa en XP; (e) Tienda de recompensas, donde los estudiantes canjean sus puntos por premios definidos por el docente; y (f) Panel de analítica, con estadísticas de progreso por alumno, por aula y —en el plan institucional— por toda la institución.')
p('El stack tecnológico emplea Angular 18 en el frontend, NestJS en el backend, PostgreSQL como base de datos relacional y se despliega en infraestructura cloud (AWS/GCP) con un modelo de alta disponibilidad y escalabilidad horizontal.')

# 1.5 Elevator Pitch
heading('1.5 Elevator Pitch', 2)
p('«¿Sabías que seis de cada diez docentes peruanos reportan baja motivación en sus alumnos todos los días? LegendaryClass convierte cada clase en una aventura RPG: los estudiantes eligen su personaje, ganan XP por participar, completan misiones y canjean recompensas reales. El docente controla todo desde su celular en tiempo real. Por menos de lo que cuesta una suscripción a Spotify, transformas la dinámica de tu aula. Empieza gratis hoy en legendaryclass.pe y descubre por qué llamamos Legendarios a nuestros docentes.»')
p('El enlace al video de presentación del pitch se encuentra disponible en el repositorio oficial del proyecto.')

page_break()

# 1.6 PESTEL
heading('1.6 Análisis PESTEL', 2)
p('El análisis PESTEL permite identificar los factores del macroentorno que influyen en la viabilidad y escalabilidad de LegendaryClass en el mercado peruano.')

add_table(
    headers=['Dimensión', 'Factor', 'Impacto', 'Tendencia'],
    rows=[
        ['Político', 'Política de transformación digital educativa (MINEDU, Plan Bicentenario 2021–2030)', 'Alto (+)', '↑ Creciente'],
        ['Político', 'Regulación de datos de menores (Ley N.° 29733 y DS 003-2013)', 'Medio (–)', '→ Estable'],
        ['Económico', 'Crecimiento del PBI peruano 2.7 % en 2025 (BCRP, 2026)', 'Medio (+)', '↑ Creciente'],
        ['Económico', 'Tipo de cambio dólar/sol (~3.75) encarece licencias de software extranjero', 'Alto (+)', '↑ Favorable'],
        ['Económico', 'Presupuesto TI limitado en colegios públicos', 'Alto (–)', '→ Estable'],
        ['Social', '7.7 millones de alumnos en EBR + 500 K en Educación Superior Tecnológica (MINEDU, 2024)', 'Alto (+)', '↑ Creciente'],
        ['Social', 'Nativos digitales como nuevo cuerpo docente (< 35 años)', 'Alto (+)', '↑ Creciente'],
        ['Tecnológico', 'Penetración de internet en docentes peruanos > 78 % (INEI, 2024)', 'Alto (+)', '↑ Creciente'],
        ['Tecnológico', 'Adopción masiva de smartphones en aulas (BYOD)', 'Alto (+)', '↑ Creciente'],
        ['Ecológico', 'Neutralidad de carbono: solución 100 % digital sin huella física', 'Bajo (+)', '→ Estable'],
        ['Legal', 'Certificación ISO/IEC 27001 recomendada para contratos con colegios', 'Medio (–)', '↑ Exigente'],
    ],
    caption='Análisis PESTEL — LegendaryClass',
    note='Elaboración propia con base en MINEDU (2024), BCRP (2026) e INEI (2024).',
    caption_num=2,
)

p('Los factores con mayor impacto positivo son la política de digitalización educativa del MINEDU, la penetración de internet en docentes y el diferencial de tipo de cambio que encarece las soluciones extranjeras. El principal riesgo es el presupuesto TI limitado de los colegios públicos, que refuerza la necesidad de un modelo de pago mensual accesible.')

page_break()

# 1.7 Porter
heading('1.7 Cinco Fuerzas de Porter', 2)
p('El modelo de Porter permite evaluar la intensidad competitiva del sector EdTech de gamificación en el Perú.')

add_table(
    headers=['Fuerza', 'Nivel', 'Principales factores'],
    rows=[
        ['Rivalidad entre competidores', 'Medio-Bajo', 'Kahoot! y ClassDojo son los competidores más cercanos, pero están enfocados en mercados anglosajones. No existe un competidor local de gamificación RPG para docentes peruanos.'],
        ['Amenaza de nuevos entrantes', 'Medio', 'Barrera tecnológica moderada (stack SaaS + gamificación). El diferenciador clave es la adaptación cultural y el precio en soles. Una startup bien financiada podría replicar el modelo en 6-12 meses.'],
        ['Poder de negociación de proveedores', 'Bajo', 'Los proveedores de infraestructura cloud (AWS, GCP, Vercel) compiten entre sí. El costo de cambio de proveedor es bajo gracias a la arquitectura containerizada.'],
        ['Poder de negociación de clientes', 'Medio-Alto', 'El docente individual tiene bajo poder; sin embargo, las instituciones (colegios, institutos) negocian licencias en bloque y pueden exigir condiciones. La sensibilidad al precio es alta.'],
        ['Amenaza de productos sustitutos', 'Medio', 'Google Classroom, Microsoft Teams for Education y grupos de WhatsApp son sustitutos funcionales parciales. No tienen mecánicas RPG pero cubren la necesidad de comunicación y gestión básica sin costo adicional.'],
    ],
    caption='Análisis de las cinco fuerzas de Porter — LegendaryClass',
    note='Elaboración propia (2026).',
    caption_num=3,
)

p('La posición competitiva de LegendaryClass es favorable: opera en un nicho (gamificación RPG para docentes peruanos) donde la rivalidad directa es baja y las barreras de entrada aumentan con el tiempo a medida que se construye la base de usuarios y el efecto red.')

# 1.8 Estrategias
heading('1.8 Estrategias Derivadas del Análisis PESTEL y Porter', 2)

heading('1.8.1 Estrategias ofensivas (aprovechar oportunidades)', 3)
estrategias_of = [
    'Alineación curricular: integrar los estándares de aprendizaje del CNEB (Currículo Nacional de Educación Básica) en las plantillas de misiones, diferenciando LegendaryClass de herramientas genéricas.',
    'Precio ancla en soles: fijar el precio base en S/. 19.90/mes/docente, por debajo del umbral perceptivo de S/. 20, aprovechando el diferencial de tipo de cambio frente a competidores en dólares.',
    'Modelo freemium para colegios públicos: ofrecer un plan gratuito con funciones básicas para penetrar el segmento público, con conversión a plan pago via institución.',
]
for e in estrategias_of:
    ep = doc.add_paragraph()
    ep.paragraph_format.first_line_indent = Cm(1.27)
    ep.paragraph_format.line_spacing = Pt(24)
    ep.paragraph_format.space_after = Pt(0)
    r = ep.add_run('• ' + e)
    r.font.name = 'Times New Roman'; r.font.size = Pt(12)

heading('1.8.2 Estrategias defensivas (mitigar amenazas)', 3)
estrategias_def = [
    'Efectos de red: diseñar funciones colaborativas entre docentes (compartir plantillas de misiones) que incrementen el costo de abandono.',
    'Cumplimiento de privacidad: obtener la certificación de protección de datos de menores antes del lanzamiento comercial para eliminar objeciones legales de colegios.',
    'Canal institucional: desarrollar un equipo de ventas B2B enfocado en contratos anuales por institución (≥ 10 docentes), que ofrezcan mayor estabilidad de ingresos y menor churn.',
]
for e in estrategias_def:
    ep = doc.add_paragraph()
    ep.paragraph_format.first_line_indent = Cm(1.27)
    ep.paragraph_format.line_spacing = Pt(24)
    ep.paragraph_format.space_after = Pt(0)
    r = ep.add_run('• ' + e)
    r.font.name = 'Times New Roman'; r.font.size = Pt(12)

page_break()

# 1.9 Justificación
heading('1.9 Justificación del Producto', 2)
p('LegendaryClass se justifica desde tres perspectivas complementarias:')
p('Desde la perspectiva pedagógica, la gamificación tiene respaldo empírico robusto. Hamari et al. (2014) revisaron 24 estudios y encontraron que el 83 % reportó mejoras en la motivación y el compromiso de los estudiantes cuando se implementaron mecánicas de juego bien diseñadas. Plass et al. (2015) demostraron que los sistemas de recompensa y los avatares personalizados incrementan la motivación intrínseca en grupos de 12 a 18 años, el rango etario principal de la EBR peruana.')
p('Desde la perspectiva de mercado, el sector EdTech latinoamericano creció un 18 % anual entre 2021 y 2024 y se proyecta alcanzar USD 3,000 millones en 2027 (HolonIQ, 2024). El Perú, con 7.7 millones de alumnos en EBR y más de 350,000 docentes activos, representa uno de los mercados más grandes de la región. La ausencia de un competidor local especializado en gamificación RPG crea una ventana de oportunidad de 18 a 24 meses para establecer liderazgo de marca.')
p('Desde la perspectiva del equipo, los integrantes combinan las capacidades técnicas, de diseño y de negocio necesarias para desarrollar, lanzar y escalar el producto. El prototipo funcional ya permite demostrar el valor de la propuesta a potenciales adoptantes tempranos.')

# 1.10 Segmentación
heading('1.10 Segmentación de Mercado', 2)
p('La segmentación se realizó mediante un modelo en cascada que parte de la población total de docentes en el Perú y aplica filtros sucesivos basados en criterios cuantitativos derivados de la encuesta y fuentes secundarias.')

add_table(
    headers=['Nivel de mercado', 'Criterio de filtro', 'Universo (docentes)', 'Tasa de filtro', 'Base de cálculo'],
    rows=[
        ['Mercado Potencial', 'Total de docentes activos en EBR y Educ. Superior Tecnológica en Perú', '350,000', '—', 'MINEDU ESCALE 2024'],
        ['Mercado Disponible', 'Docentes que trabajan actualmente (P4 = Sí)', '336,000', '96 %', 'Encuesta: 96/100 = 96 %'],
        ['Mercado Efectivo', 'Trabajan Y perciben baja motivación frecuente (P7 ≥ 4)', '188,160', '56 %', 'Encuesta: 56/100 = 56 %'],
        ['Mercado Objetivo', 'Efectivo Y dispuesto a pagar un monto (P13 ≠ No pagaría)', '169,344', '90 %', 'Encuesta: 90/100 = 90 %'],
    ],
    caption='Segmentación de mercado en cascada — LegendaryClass (2026)',
    note='Elaboración propia con base en MINEDU ESCALE (2024) y encuesta primaria (n = 100).',
    caption_num=4,
)

p('El mercado objetivo estimado asciende a 169,344 docentes, lo que representa el 48.4 % de los docentes activos en el Perú. Esta cifra constituye el universo sobre el cual se aplicarán las tasas de penetración para proyectar los ingresos del negocio.')

heading('1.10.1 Perfil del cliente ideal (Buyer Persona)', 3)
p('El cliente ideal de LegendaryClass es un docente de educación básica regular o superior técnica en el Perú, de entre 28 y 45 años, que enseña en una institución privada o privada subvencionada, cuenta con smartphone y conexión estable a internet, y ha experimentado frustración ante la baja participación de sus alumnos. Tiene disposición a probar nuevas herramientas digitales si son intuitivas y si el costo mensual no supera el equivalente a una suscripción a una plataforma de streaming.')

page_break()

# 1.11 Análisis de encuesta
heading('1.11 Análisis de la Encuesta y Fórmula Muestral', 2)

heading('1.11.1 Diseño del instrumento', 3)
p('Se diseñó una encuesta de 17 ítems distribuidos en cinco secciones: (I) Perfil del encuestado, (II) Diagnóstico del problema, (III) Interés en el producto, (IV) Disposición a pagar y (V) Condiciones de adopción. El instrumento fue validado por el equipo y revisado por el docente del curso antes de su aplicación.')

heading('1.11.2 Determinación del tamaño muestral', 3)
p('Se aplicó la fórmula para poblaciones finitas con los siguientes parámetros:')

formula_p = doc.add_paragraph()
formula_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
formula_p.paragraph_format.first_line_indent = Pt(0)
formula_p.paragraph_format.space_before = Pt(12)
formula_p.paragraph_format.space_after  = Pt(6)
formula_p.paragraph_format.line_spacing = Pt(24)
rf = formula_p.add_run('n = (Z² × p × q × N) / (e² × (N – 1) + Z² × p × q)')
rf.font.name = 'Times New Roman'; rf.font.size = Pt(12); rf.font.italic = True

params_p = doc.add_paragraph()
params_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
params_p.paragraph_format.first_line_indent = Pt(0)
params_p.paragraph_format.space_after = Pt(6)
params_p.paragraph_format.line_spacing = Pt(24)
rp = params_p.add_run('Z = 1.96 (95 % de confianza),  p = q = 0.50,  e = 9.76 %,  N = 350,000')
rp.font.name = 'Times New Roman'; rp.font.size = Pt(12)

result_p = doc.add_paragraph()
result_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
result_p.paragraph_format.first_line_indent = Pt(0)
result_p.paragraph_format.space_after = Pt(12)
result_p.paragraph_format.line_spacing = Pt(24)
rr = result_p.add_run('n = (3.8416 × 0.25 × 350,000) / (0.009527 × 349,999 + 3.8416 × 0.25) ≈ 100')
rr.font.name = 'Times New Roman'; rr.font.size = Pt(12)

p('El tamaño muestral calculado es de 100 encuestas, lo que proporciona un margen de error del ± 9.76 % para una población de 350,000 docentes con un nivel de confianza del 95 %. La muestra fue seleccionada de manera no probabilística por conveniencia entre docentes activos con acceso a redes sociales y aplicaciones de mensajería.')

heading('1.11.3 Resultados de P7 — Percepción de baja motivación', 3)
p('La pregunta P7 evalúa la frecuencia con que el docente percibe baja motivación en sus estudiantes, en una escala de Likert de 1 (Nunca) a 5 (Siempre). Los resultados se presentan en la Tabla 5.')

add_table(
    headers=['Respuesta', 'Frecuencia (n)', 'Porcentaje (%)'],
    rows=[
        ['1 — Nunca', '3', '3.0 %'],
        ['2 — Raramente', '11', '11.0 %'],
        ['3 — A veces', '30', '30.0 %'],
        ['4 — Frecuentemente', '36', '36.0 %'],
        ['5 — Siempre', '20', '20.0 %'],
        ['Total', '100', '100.0 %'],
    ],
    caption='Distribución de respuestas — P7: Frecuencia de baja motivación estudiantil',
    note='Elaboración propia a partir de la encuesta de preferencia (n = 100, 2026).',
    caption_num=5,
)

p('El 56 % de los encuestados reporta percibir baja motivación de manera frecuente o constante (categorías 4 y 5), lo que confirma la prevalencia del problema central que LegendaryClass busca resolver. Este porcentaje se utiliza como tasa de filtro para determinar el Mercado Efectivo.')

heading('1.11.4 Resultados de P13 — Disposición a pagar', 3)
p('La pregunta P13 indaga el monto mensual que el docente estaría dispuesto a pagar por acceder a la plataforma completa. Es la variable de mayor relevancia para la modelización financiera.')

add_table(
    headers=['Rango de precio mensual', 'Frecuencia (n)', 'Porcentaje (%)', 'Punto medio (S/.)'],
    rows=[
        ['Menos de S/. 15', '45', '50.0 %', '10.00'],
        ['Entre S/. 15 y S/. 25', '28', '31.1 %', '20.00'],
        ['Entre S/. 26 y S/. 40', '15', '16.7 %', '33.00'],
        ['Más de S/. 40', '2', '2.2 %', '45.00'],
        ['No pagaría (excluido)', '10', '—', '—'],
        ['Subtotal (respondentes válidos)', '90', '100.0 %', '—'],
    ],
    caption='Distribución de respuestas — P13: Disposición a pagar mensual',
    note='Se excluyen 10 encuestados que respondieron "No pagaría por ningún monto". Elaboración propia (2026).',
    caption_num=6,
)

heading('1.11.5 Cálculo del precio ponderado', 3)
p('El precio mensual ponderado se calcula como la media aritmética ponderada de los puntos medios de cada rango, usando como pesos las frecuencias relativas de los respondentes válidos (n = 90):')

add_table(
    headers=['Rango', 'Peso (f_i / n)', 'Punto medio (m_i)', 'Contribución (f_i × m_i / n)'],
    rows=[
        ['< S/. 15', '0.5000', 'S/. 10.00', 'S/. 5.00'],
        ['S/. 15–25', '0.3111', 'S/. 20.00', 'S/. 6.22'],
        ['S/. 26–40', '0.1667', 'S/. 33.00', 'S/. 5.50'],
        ['> S/. 40',  '0.0222', 'S/. 45.00', 'S/. 1.00'],
        ['Precio mensual ponderado', '—', '—', 'S/. 17.72'],
        ['Precio anual ponderado (× 12)', '—', '—', 'S/. 212.64'],
    ],
    caption='Cálculo del precio ponderado mensual — LegendaryClass',
    note='Elaboración propia. El precio ponderado de S/. 17.72/mes refleja la baja capacidad de pago individual del docente peruano. Referencia: ClassDojo Plus (~S/. 25/mes), Kahoot! Pro (~S/. 37/mes).',
    caption_num=7,
)

p('El precio ponderado de S/. 17.72/mes (S/. 212.64/año) es el más bajo dentro del rango de mercado comparado, lo que refleja con fidelidad la capacidad de pago real del docente peruano promedio. Esta cifra, inferior a ClassDojo Plus (≈ S/. 25/mes) y Kahoot! Pro (≈ S/. 37/mes), refuerza la estrategia de penetración por precio y justifica el modelo freemium como palanca de adquisición inicial.')

page_break()

# 1.12 Validación
heading('1.12 Validación del Mercado', 2)
p('La validación del mercado se realizó mediante tres mecanismos complementarios:')

p('Encuesta primaria (n = 100). El 72 % de los encuestados que trabajan actualmente (P4) respondió afirmativamente a P9 («¿Estaría interesado en usar LegendaryClass?»), distribuyéndose entre «Sí, definitivamente» (31 %) y «Tal vez, necesito ver más detalles» (41 %). Solo el 27 % respondió «No». Este nivel de interés declarado (72 %) es superior al umbral de viabilidad del 50 % utilizado en estudios de intención de compra para productos SaaS en etapa de concepto (HolonIQ, 2024).')

p('Pruebas de concepto con docentes reales. Se realizaron cinco sesiones de demostración con docentes de instituciones educativas de Lima Metropolitana (tres privadas, dos públicas). En promedio, los participantes calificaron la propuesta de valor con 8.2/10 en escala de utilidad percibida y señalaron como principales objeciones: (a) requieren ver que funcione con alumnos reales antes de pagar, y (b) preferirían que la institución asuma el costo. Estas observaciones orientaron el diseño del modelo freemium y del canal institucional.')

p('Benchmarking competitivo. Se analizaron tres herramientas comparables: Kahoot! (quiz gamificado, S/. 37/mes), ClassDojo (gestión conductual, S/. 25/mes) y Duolingo for Schools (idiomas, gratuito con monetización de datos). Ninguna ofrece el sistema de personajes RPG, la tienda de recompensas configurable por el docente ni la integración completa de XP, misiones y comportamientos en un solo panel. Esto confirma la diferenciación del producto.')

# 1.13 Demanda e ingresos
heading('1.13 Estimación de la Demanda e Ingresos', 2)
p('Con base en el mercado objetivo de 169,344 docentes y el precio ponderado anual de S/. 212.64, se proyectan tres escenarios para el primer año de operación (2027), diferenciados por la tasa de penetración aplicada:')

add_table(
    headers=['Escenario', 'Tasa de penetración', 'Suscriptores Y1', 'Precio anual (S/.)', 'Ingresos Y1 (S/.)'],
    rows=[
        ['Optimista', '1.5 %', '2,540', '212.64', '540,106'],
        ['Moderado',  '0.7 %', '1,185', '212.64', '251,978'],
        ['Pesimista', '0.2 %',   '338', '212.64',  '71,872'],
    ],
    caption='Escenarios de demanda — Año 1 (2027)',
    note='La tasa de penetración del 0.7 % para el escenario moderado refleja el crecimiento orgánico sin inversión publicitaria significativa, coherente con el desempeño de startups EdTech en etapa temprana en Latinoamérica (HolonIQ, 2024). Elaboración propia.',
    caption_num=8,
)

p('El escenario moderado (0.7 % de penetración, 1,185 suscriptores, S/. 251,978 en ingresos) es el más plausible considerando que: (a) el equipo no cuenta con capital inicial para publicidad masiva, (b) el producto aún no tiene tracción de mercado comprobada y (c) el ciclo de ventas a instituciones educativas puede extenderse de tres a seis meses. El escenario optimista (1.5 %) sería alcanzable si se logra al menos un contrato institucional con 150–200 docentes en el primer año.')

p('Nota: Las cifras anteriores contemplan únicamente el plan de pago individual por docente (S/. 212.64/año). El plan institucional (licencia anual para toda la institución) se negocia en términos B2B y no se incluye en esta proyección base para mantener el conservadurismo del modelo.')

page_break()

# 1.14 Proyección 5 años
heading('1.14 Proyección Financiera a Cinco Años (2027–2031)', 2)
p('La proyección financiera aplica un CAGR del 5 % anual sobre el número de suscriptores para los tres escenarios, basándose en el extremo conservador del rango 5–8 % reportado por Statista y HolonIQ (2024) para el mercado EdTech SaaS en países emergentes de Latinoamérica. El precio base se mantiene constante durante el período proyectado para preservar el conservadurismo del modelo.')

add_table(
    headers=['Año', 'Suscriptores', 'Precio anual (S/.)', 'Ingresos (S/.)', 'Crecimiento YoY'],
    rows=[
        ['2027 (Y1)', '1,185', '212.64', '251,978', '—'],
        ['2028 (Y2)', '1,244', '212.64', '264,524', '5.0 %'],
        ['2029 (Y3)', '1,306', '212.64', '277,708', '5.0 %'],
        ['2030 (Y4)', '1,371', '212.64', '291,529', '5.0 %'],
        ['2031 (Y5)', '1,440', '212.64', '306,202', '5.0 %'],
        ['Acumulado 5 años', '—', '—', '1,391,941', '—'],
    ],
    caption='Proyección financiera 2027–2031 — Escenario moderado (plan individual)',
    note='CAGR de suscriptores: 5 % anual (punto inferior del rango EdTech LATAM). Precio constante S/. 212.64/año. No incluye ingresos por plan institucional. Elaboración propia con base en HolonIQ (2024) y Statista (2024).',
    caption_num=9,
)

p('El modelo proyecta ingresos acumulados de S/. 1.39 millones en cinco años bajo el escenario moderado, con un punto de equilibrio estimado entre el segundo y tercer año asumiendo costos fijos de operación de S/. 120,000/año (infraestructura cloud, mantenimiento del equipo y gastos básicos de marketing digital). El escenario optimista (1.5 %, CAGR 5 %) genera ingresos acumulados de aproximadamente S/. 2.97 millones en el mismo período.')

p('Estas cifras son deliberadamente conservadoras. La ausencia de capital inicial de marketing y el ciclo de ventas largo en el sector educativo justifican asumir la tasa mínima de crecimiento del segmento. Cualquier contrato institucional adicional o ronda seed elevaría sustancialmente estas proyecciones sin requerir cambios en el modelo de negocio base.')

page_break()

# CONCLUSIONES
heading('Conclusiones', 1)
conclusiones = [
    'LegendaryClass responde a un problema educativo real y prevalente: el 56 % de los docentes encuestados reporta baja motivación estudiantil de manera frecuente o constante, lo que valida la hipótesis central del proyecto.',
    'El mercado objetivo en el Perú asciende a 169,344 docentes, estimados a partir de la aplicación de filtros en cascada sobre una base de 350,000 docentes activos en EBR y Educación Superior Tecnológica. El 90 % de los encuestados con el problema percibido declara disposición a pagar algún monto.',
    'El precio mensual ponderado de S/. 17.72 (S/. 212.64/año) refleja con fidelidad la baja capacidad de pago individual del docente peruano y se posiciona por debajo de los benchmarks regionales (ClassDojo Plus ≈ S/. 25/mes, Kahoot! Pro ≈ S/. 37/mes), lo que refuerza la propuesta de valor diferenciada.',
    'Bajo un escenario moderado (0.7 % de penetración, CAGR 5 %), el proyecto proyecta ingresos de S/. 251,978 en el primer año y S/. 1.39 millones acumulados en cinco años, cifras conservadoras y alcanzables para una startup EdTech sin capital inicial de marketing.',
    'El análisis PESTEL y de las cinco fuerzas de Porter confirma que LegendaryClass opera en un entorno favorable: política de digitalización educativa en expansión, bajo nivel de rivalidad competitiva local y diferencial de tipo de cambio favorable frente a competidores en dólares.',
    'El principal riesgo identificado es la alta sensibilidad al precio del segmento docente y la dependencia de decisiones institucionales para la adopción masiva, lo que refuerza la necesidad de un modelo freemium y un canal de ventas B2B.',
]
for i, c in enumerate(conclusiones, 1):
    cp = doc.add_paragraph()
    cp.paragraph_format.first_line_indent = Cm(1.27)
    cp.paragraph_format.line_spacing = Pt(24)
    cp.paragraph_format.space_after = Pt(0)
    r1 = cp.add_run(f'{i}. ')
    r1.font.name = 'Times New Roman'; r1.font.size = Pt(12); r1.font.bold = True
    r2 = cp.add_run(c)
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(12)

page_break()

# RECOMENDACIONES
heading('Recomendaciones', 1)
recomendaciones = [
    'Ampliar la muestra de la encuesta a un mínimo de 384 respondentes (margen de error ±5 % para N = 350,000) antes del lanzamiento oficial, utilizando paneles de investigación en línea para garantizar representatividad geográfica (Lima, provincias) y por tipo de institución (pública/privada).',
    'Desarrollar el plan institucional (licencia anual por colegio) como producto estrella para el canal B2B. Un contrato de 50 docentes a S/. 250/docente/año genera S/. 12,500 de ARR (Annual Recurring Revenue) y reduce el costo de adquisición por suscriptor hasta en un 70 % respecto al canal individual.',
    'Implementar el modelo freemium (plan gratuito con límite de 1 aula y 30 alumnos) antes del lanzamiento pago, para construir base de usuarios y generar testimonios de docentes que validen la propuesta de valor.',
    'Certificar el cumplimiento de la Ley N.° 29733 (Ley de Protección de Datos Personales del Perú) y publicar la política de privacidad antes de cualquier acuerdo con instituciones educativas públicas.',
    'Explorar alianzas estratégicas con editoriales educativas (Santillana, SM, Norma) y distribuidores de materiales escolares para crear un canal de distribución indirecto de bajo costo.',
    'Postular al Fondo CONCYTEC-PROCIENCIA de emprendimiento tecnológico 2026 para obtener financiamiento semilla no dilutivo que permita contratar al menos un desarrollador adicional y lanzar campañas de marketing digital.',
]
for i, r_text in enumerate(recomendaciones, 1):
    rp_p = doc.add_paragraph()
    rp_p.paragraph_format.first_line_indent = Cm(1.27)
    rp_p.paragraph_format.line_spacing = Pt(24)
    rp_p.paragraph_format.space_after = Pt(0)
    r1 = rp_p.add_run(f'{i}. ')
    r1.font.name = 'Times New Roman'; r1.font.size = Pt(12); r1.font.bold = True
    r2 = rp_p.add_run(r_text)
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(12)

page_break()

# REFERENCIAS
heading('Referencias', 1)

refs = [
    'Banco Central de Reserva del Perú. (2026). Reporte de inflación: Panorama actual y proyecciones macroeconómicas 2026–2027. BCRP. https://www.bcrp.gob.pe/publicaciones/reporte-de-inflacion.html',
    'Deterding, S., Dixon, D., Khaled, R., & Nacke, L. (2011). From game design elements to gamefulness: Defining "gamification." En Proceedings of the 15th International Academic MindTrek Conference (pp. 9–15). ACM. https://doi.org/10.1145/2181037.2181040',
    'Hamari, J., Koivisto, J., & Sarsa, H. (2014). Does gamification work? — A literature review of empirical studies on gamification. En Proceedings of the 47th Hawaii International Conference on System Sciences (pp. 3025–3034). IEEE. https://doi.org/10.1109/HICSS.2014.377',
    'HolonIQ. (2024). Latin America EdTech market report 2024. HolonIQ Intelligence. https://www.holoniq.com/notes/latin-america-edtech-market-2024',
    'Instituto Nacional de Estadística e Informática. (2024). Encuesta nacional de hogares sobre condiciones de vida y pobreza 2023 — Módulo de tecnologías de información y comunicación. INEI. https://www.inei.gob.pe/estadisticas/indice-tematico/tecnologias-de-la-informacion-y-telecomunicaciones/',
    'Ministerio de Educación del Perú. (2024). Estadísticas de la calidad educativa (ESCALE) — Magnitudes de la educación básica regular 2024. MINEDU. http://escale.minedu.gob.pe/',
    'Plass, J. L., Homer, B. D., & Kinzer, C. K. (2015). Foundations of game-based learning. Educational Psychologist, 50(4), 258–283. https://doi.org/10.1080/00461520.2015.1122533',
]

for ref in refs:
    ref_p = doc.add_paragraph()
    ref_p.paragraph_format.first_line_indent = Cm(-1.27)   # Sangría francesa
    ref_p.paragraph_format.left_indent = Cm(1.27)
    ref_p.paragraph_format.line_spacing = Pt(24)
    ref_p.paragraph_format.space_after = Pt(0)
    r = ref_p.add_run(ref)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

# ─────────────────────────────────────────────────────────────
# GUARDAR
# ─────────────────────────────────────────────────────────────
output_path = 'Informe_LaboratorioPractica1_LegendaryClass.docx'
doc.save(output_path)
print(f'Informe generado: {output_path}')
print(f'Páginas estimadas: ~32')
